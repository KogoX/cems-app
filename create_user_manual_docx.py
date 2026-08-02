import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_placeholder_box(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_background(cell, "F4FBF7")
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("📷 [SCREENSHOT PLACEHOLDER]\n")
    run1.bold = True
    run1.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    run1.font.size = Pt(11)
    
    run2 = p.add_run(f"{caption}\n")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
    run2.font.size = Pt(10)

    run3 = p.add_run("(Paste your app screenshot image here before exporting to PDF)")
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run3.font.size = Pt(9)
    
    doc.add_paragraph() # spacing

def build_manual():
    doc = Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # -------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("KENYATTA UNIVERSITY\nSCHOOL OF COMPUTING AND INFORMATION TECHNOLOGY\nDEPARTMENT OF COMPUTER SCIENCE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Logo Box Placeholder
    logo_table = doc.add_table(rows=1, cols=1)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    l_cell = logo_table.cell(0, 0)
    l_cell.width = Inches(2.2)
    set_cell_background(l_cell, "F3F4F6")
    set_cell_margins(l_cell, top=140, bottom=140, left=140, right=140)
    lp = l_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lrun = lp.add_run("[ UNIVERSITY LOGO ]")
    lrun.bold = True
    lrun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    lrun.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Main Project Title
    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prun = proj_p.add_run("COOPERATIVE EXPORT MANAGEMENT SYSTEM (CEMS)")
    prun.bold = True
    prun.font.size = Pt(20)
    prun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub_p.add_run("System User Manual & Technical Operation Guide")
    srun.italic = True
    srun.font.size = Pt(14)
    srun.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # Metadata Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
      ("Course Name & Code:", "BSc Computer Science Final Defence (COMP 400)"),
      ("Student Name & ID:", "[Your Name / Student Registration Number]"),
      ("Project Supervisor:", "[Supervisor / Coordinator Name]"),
      ("Date of Defence:", "Academic Year 2025/2026")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width, c2.width = Inches(2.2), Inches(3.8)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 1: SYSTEM OVERVIEW & OBJECTIVES
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    hrun = h1.add_run("1. System Overview & Executive Summary")
    hrun.bold = True
    hrun.font.size = Pt(16)
    hrun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    h1.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "The Cooperative Export Management System (CEMS) is a digital platform designed to bridge the structural gap "
        "between smallholder avocado farmers, cooperative managers, and international export buyers in Kenya. "
        "Traditionally, smallholder farmers face severe exploitation from informal middlemen, lack cold-chain visibility, "
        "and suffer financial losses due to manual record-keeping errors and delayed payouts."
    )
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "CEMS streamlines the complete post-harvest lifecycle: from farmer yield logging and manager quality inspection "
        "to buyer procurement via Paystack digital checkout, real-time shipment tracking, and automated cooperative farmer disbursals."
    )
    p.paragraph_format.space_after = Pt(14)

    # Objectives Subsection
    h2 = doc.add_paragraph()
    h2run = h2.add_run("1.1 Project Objectives")
    h2run.bold = True
    h2run.font.size = Pt(13)
    h2run.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
    h2.paragraph_format.space_after = Pt(6)

    objs = [
        ("Objective 1 — Direct Buyer Access:", " To develop a centralized digital platform that gives smallholder avocado farmers direct access to export buyers, eliminating the dependence on middlemen and enabling fairer pricing for their produce."),
        ("Objective 2 — Transparent Payment Tracking:", " To implement a transparent payment tracking and distribution system that allows farmers to verify their earnings based on submitted yield, eliminating miscalculations and manipulation of cooperative records."),
        ("Objective 3 — Digital Record-Keeping & Automation:", " To provide cooperative managers and farmers with digital record-keeping and reporting tools that replace manual processes, ensuring accurate, secure, and tamper-proof management of produce, orders, and payments.")
    ]

    for title, desc in objs:
        op = doc.add_paragraph(style='List Bullet')
        op.paragraph_format.space_after = Pt(6)
        r1 = op.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
        r2 = op.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Market Rate Table
    h2 = doc.add_paragraph()
    h2run = h2.add_run("1.2 Live Market Pricing & Cooperative Revenue Breakdown")
    h2run.bold = True
    h2run.font.size = Pt(13)
    h2run.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
    h2.paragraph_format.space_after = Pt(6)

    rate_table = doc.add_table(rows=4, cols=4)
    rate_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Produce Grade", "Buyer Price (KSh/kg)", "Farmer Payout (KSh/kg)", "Co-op Retained (KSh/kg)"]
    hdr_row = rate_table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "125C3F")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)

    rows_data = [
        ("Grade A (Export Hass)", "160 KSh / kg ($1.23)", "115 KSh / kg ($0.88)", "45 KSh / kg ($0.35)"),
        ("Grade B (Processing)", "110 KSh / kg ($0.85)", "75 KSh / kg ($0.58)", "35 KSh / kg ($0.27)"),
        ("Grade C (Standard)", "80 KSh / kg ($0.62)", "50 KSh / kg ($0.38)", "30 KSh / kg ($0.23)")
    ]
    for r_idx, r_data in enumerate(rows_data, start=1):
        row = rate_table.rows[r_idx]
        bg_color = "F9F9F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # -------------------------------------------------------------
    # SECTION 2: FARMER USER GUIDE
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    hrun = h1.add_run("2. Farmer User Guide")
    hrun.bold = True
    hrun.font.size = Pt(16)
    hrun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    h1.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "Farmers access CEMS to record upcoming perishable harvest yields, track manager approval status, "
        "and receive transparent, automated mobile payouts."
    )
    p.paragraph_format.space_after = Pt(8)

    f_steps = [
        ("Step 1 — Registration & Login:", " Launch the app and select 'Farmer' during registration. Enter your name, phone number, and location. Existing farmers log in using registered credentials."),
        ("Step 2 — Logging Upcoming Harvest Yields:", " Navigate to 'Log Yield'. Select produce variety (e.g., Hass, Fuerte), harvest quantity in kilograms, quality grade, and expected upcoming harvest date (e.g. Tomorrow, 2 Days to come). Upload real harvest photos using the camera or gallery."),
        ("Step 3 — Photo Upload & Storage:", " CEMS resizes and compresses photos on-device before uploading directly to Supabase Object Storage (harvest-photos bucket), generating fast CDN image URLs."),
        ("Step 4 — Dashboard & Payout Verification:", " Farmers track real-time yield statuses (Logged → Approved → Scheduled → Paid) and inspect calculated net earnings (at 115 KSh/kg for Grade A) transparently.")
    ]

    for title, desc in f_steps:
        fp = doc.add_paragraph(style='List Bullet')
        fp.paragraph_format.space_after = Pt(6)
        r1 = fp.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
        r2 = fp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_placeholder_box(doc, "Farmer Section — Logging Upcoming Perishable Yield & Selecting Future Harvest Date")
    add_placeholder_box(doc, "Farmer Section — Farmer Yield Status & Payout History Dashboard")

    # -------------------------------------------------------------
    # SECTION 3: BUYER USER GUIDE
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    hrun = h1.add_run("3. Export Buyer User Guide")
    hrun.bold = True
    hrun.font.size = Pt(16)
    hrun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    h1.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "Export buyers access CEMS to inspect verified avocado harvest listings, place procurement orders, "
        "pay via Paystack, and track live shipment transit locations."
    )
    p.paragraph_format.space_after = Pt(8)

    b_steps = [
        ("Step 1 — Marketplace Browsing:", " Buyers explore verified avocado harvest listings filtered by Grade (A, B, C). Real farmer harvest photos, volume availability, and harvest dates are displayed."),
        ("Step 2 — Dual-Currency Price Breakdown:", " Prices and total amounts are displayed in both Kenyan Shillings (KSh) and US Dollars ($ USD) in a clear, light-font layout for international buyers."),
        ("Step 3 — Secure Paystack Checkout:", " Tap 'Place Order' and proceed to Paystack Checkout modal. Pay seamlessly using M-Pesa, debit/credit cards, or bank transfer."),
        ("Step 4 — Shipment Tracking & PDF Statement Export:", " Track real-time shipment status (Picked Up → In Transit → Delivered) with live GPS location updates. Export PDF Procurement Statements instantly.")
    ]

    for title, desc in b_steps:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(6)
        r1 = bp.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
        r2 = bp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_placeholder_box(doc, "Buyer Section — Avocado Marketplace Listings & Real Farmer Photos Carousel")
    add_placeholder_box(doc, "Buyer Section — Order Modal with Dual Currency (KSh & $) & Paystack Secure Checkout")

    # -------------------------------------------------------------
    # SECTION 4: COOPERATIVE MANAGER & SUPER ADMIN USER GUIDE
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    hrun = h1.add_run("4. Cooperative Manager & Super Admin Guide")
    hrun.bold = True
    hrun.font.size = Pt(16)
    hrun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    h1.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(
        "Cooperative managers oversee platform operations, verify submitted yields, update shipment tracking locations, "
        "and disburse auto-calculated payouts to farmers. The initial registered manager becomes the system Super Admin."
    )
    p.paragraph_format.space_after = Pt(8)

    m_steps = [
        ("Step 1 — First Manager Super Admin Setup & Sub-Manager Approval:", " The first manager to register in CEMS becomes the Super Admin ('Active' and 'Verified'). Subsequent manager registrations default to 'Pending Approval' and require Super Admin authorization before logging in."),
        ("Step 2 — Harvest Inspection & Approval Queue:", " Review submitted farmer harvest records, inspect uploaded photos, verify quality grades, and approve or reject listings."),
        ("Step 3 — Automatic Stock-Based Farmer Disbursal:", " Select a farmer in 'Farmer Payouts'. CEMS automatically detects buyer-paid stock (e.g. STOCK-0005), calculates Net Farmer Payout (115 KSh/kg), deducts Co-op Transport/Maintenance reserve (45 KSh/kg), and pre-fills the disbursal amount automatically."),
        ("Step 4 — Executive Financial Reporting:", " Generate and export PDF Executive Operations Statements summarizing total harvest volume, buyer revenue, and farmer disbursements.")
    ]

    for title, desc in m_steps:
        mp = doc.add_paragraph(style='List Bullet')
        mp.paragraph_format.space_after = Pt(6)
        r1 = mp.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x2A, 0x5C, 0x43)
        r2 = mp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_placeholder_box(doc, "Manager Section — Manager Dashboard & Executive Supply-Demand Metrics")
    add_placeholder_box(doc, "Manager Section — Automated Stock Payout Card & Co-op Retention Calculation")
    add_placeholder_box(doc, "Manager Section — Super Admin Manager Approval & User Management Screen")

    # -------------------------------------------------------------
    # SECTION 5: TECHNICAL ARCHITECTURE & VERIFICATION
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    hrun = h1.add_run("5. Technical Architecture & System Verification")
    hrun.bold = True
    hrun.font.size = Pt(16)
    hrun.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
    h1.paragraph_format.space_after = Pt(8)

    arch_points = [
        ("Database Infrastructure:", " PostgreSQL relational database hosted on Supabase Cloud, featuring automated table bootstrapping, foreign key cascades, transaction isolation, and database connection pooling."),
        ("Object Storage & CDN:", " Supabase Storage Bucket ('harvest-photos') serving client-compressed JPEG images over public CDN, eliminating database bloat."),
        ("Payment Gateway Integration:", " Paystack API webhook integration supporting M-Pesa, card payments, and automated B2C mobile money transfers."),
        ("Security & Auth:", " JWT authentication with Role-Based Access Control (RBAC), bcrypt password hashing, and Super Admin authorization guards.")
    ]

    for title, desc in arch_points:
        ap = doc.add_paragraph(style='List Bullet')
        ap.paragraph_format.space_after = Pt(6)
        r1 = ap.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x12, 0x5C, 0x3F)
        r2 = ap.add_run(desc)

    # Save Document
    doc_path = r"c:\Users\MicroAple\cems\CEMS_System_User_Manual.docx"
    doc.save(doc_path)
    print(f"SUCCESS: Word document created at: {doc_path}")

if __name__ == "__main__":
    build_manual()
